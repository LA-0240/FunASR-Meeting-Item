from django.conf import settings
from django.http import FileResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.views import View
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from openai import OpenAI
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import time
import uuid
from datetime import datetime
import traceback
import numpy as np
from .utils import extract_audio_from_video
from .models import Voiceprint, asr_model
from .voiceprint_utils import extract_voiceprint_feature, check_voiceprint_duplicate, match_voiceprint


# ------------------- 视频上传+语音转写 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class VideoASRTranscribeView(APIView):
    def post(self, request):
        """
        视频上传接口：提取音频→ASR转写→说话人分离→声纹匹配
        返回：带正确说话人名的转写文本
        """
        temp_video = None
        temp_audio = None
        try:
            # 1. 校验文件是否上传
            if 'file' not in request.FILES:
                return Response(
                    {"status": "failed", "detail": "未上传视频文件"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            file = request.FILES['file']
            # 2. 校验视频格式
            if not file.name.lower().endswith(settings.ALLOWED_VIDEO_EXTENSIONS):
                return Response(
                    {"status": "failed", "detail": f"仅支持以下视频格式：{settings.ALLOWED_VIDEO_EXTENSIONS}"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 保存临时视频文件
            video_filename = f"temp_video_{datetime.now().strftime('%Y%m%d%H%M%S_%f')}_{file.name}"
            temp_video = os.path.join(settings.TEMP_DIR, video_filename)
            with open(temp_video, "wb") as f:
                for chunk in file.chunks():
                    f.write(chunk)
            
            # 4. 提取音频（WAV格式）
            audio_filename = f"temp_audio_{uuid.uuid4()}.wav"
            temp_audio = os.path.join(settings.TEMP_DIR, audio_filename)
            if not extract_audio_from_video(temp_video, temp_audio):
                return Response(
                    {"status": "failed", "detail": "视频音频提取失败"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # 5. 解析请求参数
            batch_size_s = request.POST.get("batch_size_s", "")
            try:
                batch_size_s = int(batch_size_s.strip()) if batch_size_s.strip() else 300
            except (ValueError, TypeError):
                batch_size_s = 300
            hotword = request.POST.get("hotword", None)
            
            # 6. 调用ASR模型
            print(f"[{datetime.now()}] 视频处理第一步：说话人分离 + ASR识别")
            try:
                result = asr_model.generate(
                    input=temp_audio,
                    batch_size_s=batch_size_s,
                    hotword=hotword,
                    punc=True,
                    spk_segment=True,
                    speaker_diarization=True,
                    merge_vad=True,
                    max_single_segment_time=30
                )
            except IndexError as e:
                print(f"{datetime.now()} - 说话人识别失败，降级为纯文本转写：{str(e)}")
                result = asr_model.generate(
                    input=temp_audio,
                    batch_size_s=batch_size_s,
                    hotword=hotword,
                    punc=True,
                    spk_segment=False,
                    merge_vad=True,
                    max_single_segment_time=30
                )
            except Exception as e:
                return Response(
                    {"status": "failed", "detail": f"ASR模型处理失败：{str(e)}"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            sentence_info = result[0].get("sentence_info", []) if result else []
            if not sentence_info:
                return Response({
                    "status": "success",
                    "filename": file.name,
                    "transcription": [{
                        "spk": "未知说话人",
                        "text": "未识别到有效语音内容",
                        "start_time": 0.00,
                        "end_time": 0.00,
                        "original_spk": 0
                    }],
                    "speaker_stats": {"total_speakers":0,"speaker_ids":[],"matched_speakers":{}},
                    "note": "视频处理完成",
                    "timestamp": datetime.now().isoformat()
                })

            # ===================== 收集说话人时间段 =====================
            from collections import defaultdict
            import subprocess
            speaker_segments = defaultdict(list)
            speaker_ids = set()

            for seg in sentence_info:
                spk_id = seg.get("spk") or seg.get("sp") or 0
                start_ms = seg.get("start", 0)
                end_ms = seg.get("end", 0)
                speaker_segments[spk_id].append((start_ms, end_ms))
                speaker_ids.add(spk_id)

            # ===================== 第二步：每人独立声纹匹配（核心） =====================
            print(f"[{datetime.now()}] 视频处理第二步：{len(speaker_segments)} 个说话人进行声纹匹配")
            speaker_name_map = {}
            from .voiceprint_utils import extract_voiceprint_feature, match_voiceprint

            for spk_id, segments in speaker_segments.items():
                segments.sort(key=lambda x: x[1]-x[0], reverse=True)
                best_start, best_end = segments[0]
                start_sec = best_start / 1000.0
                duration_sec = (best_end - best_start) / 1000.0
                clip_path = os.path.join(settings.TEMP_DIR, f"vid_clip_{spk_id}.wav")

                # ffmpeg 切割（无pydub）
                subprocess.run([
                    "ffmpeg",
                    "-ss", str(start_sec),
                    "-t", str(duration_sec),
                    "-i", temp_audio,
                    "-ar", "16000",
                    "-ac", "1",
                    "-y",
                    clip_path
                ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

                # 声纹识别
                feat = extract_voiceprint_feature(clip_path)
                if feat is not None:
                    name = match_voiceprint(feat)
                    speaker_name_map[spk_id] = name if name else f"spk-{spk_id}"
                else:
                    speaker_name_map[spk_id] = f"spk-{spk_id}"

                if os.path.exists(clip_path):
                    os.remove(clip_path)

            # ===================== 最终结果拼接 =====================
            formatted_result = []
            matched_speakers = {}
            for spk_id, name in speaker_name_map.items():
                if not name.startswith("spk-"):
                    matched_speakers[spk_id] = name

            for seg in sentence_info:
                spk_id = seg.get("spk") or seg.get("sp") or 0
                name = speaker_name_map.get(spk_id, f"spk-{spk_id}")
                formatted_result.append({
                    "spk": name,
                    "text": seg.get("text", "").strip(),
                    "start_time": round(seg.get("start", 0)/1000, 2),
                    "end_time": round(seg.get("end", 0)/1000, 2),
                    "original_spk": spk_id
                })

            # ===================== 返回 =====================
            return Response({
                "status": "success",
                "filename": file.name,
                "transcription": formatted_result,
                "speaker_stats": {
                    "total_speakers": len(speaker_ids),
                    "speaker_ids": sorted(list(speaker_ids)),
                    "matched_speakers": matched_speakers
                },
                "note": "视频已处理：分离+声纹识别完成",
                "timestamp": datetime.now().isoformat()
            })

        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"视频处理失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        finally:
            for temp_file in [temp_video, temp_audio]:
                if temp_file and os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                    except:
                        pass

# 初始化LLM客户端
llm_client = OpenAI(
    api_key=settings.LLM_CONFIG["api_key"],
    base_url=settings.LLM_CONFIG["base_url"]
)

# ------------------- 基础接口（健康检查） -------------------
class HealthCheckView(APIView):
    def get(self, request):
        """服务健康检查，对齐FastAPI的/接口"""
        return Response({
            "status": "healthy",
            "service": "基于LLM的会议纪要智能生成系统",
            "timestamp": datetime.now().isoformat(),
            "features": ["语音|视频---转文字（带时间标点/说话人）", "会议纪要生成","会议摘要生成","Word文档导出(转写、纪要、摘要)"]
        })

# ------------------- ASR语音转文字接口 -------------------
# ------------------- ASR语音转文字接口 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class ASRTranscribeView(APIView):
    def post(self, request):
        """语音转文字接口，支持多说话人声纹识别（两步法：分离→切割→匹配）"""
        temp_file = None
        try:
            # 1. 校验文件上传
            if 'file' not in request.FILES:
                return Response(
                    {"status": "failed", "detail": "未上传音频文件"},
                    status=status.HTTP_400_BAD_REQUEST
                )

            file = request.FILES['file']
            if not file.name.lower().endswith(settings.ALLOWED_EXTENSIONS):
                return Response(
                    {"status": "failed", "detail": f"仅支持格式：{settings.ALLOWED_EXTENSIONS}"},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # 2. 保存临时文件
            temp_filename = f"temp_{datetime.now().strftime('%Y%m%d%H%M%S_%f')}_{file.name}"
            temp_file = os.path.join(settings.TEMP_DIR, temp_filename)
            with open(temp_file, 'wb') as f:
                for chunk in file.chunks():
                    f.write(chunk)

            # 3. 解析参数
            batch_size_s = int(request.POST.get("batch_size_s", 300))
            hotword = request.POST.get("hotword", None)

            # ===================== 步骤1：FunASR 说话人分离 + ASR识别 =====================
            print(f"[{datetime.now()}] 第一步：说话人分离 + 语音识别")
            result = asr_model.generate(
                input=temp_file,
                batch_size_s=batch_size_s,
                hotword=hotword,
                punc=True,
                spk_segment=True,
                speaker_diarization=True,
                merge_vad=True,
                max_single_segment_time=30
            )

            # 4. 格式化结果
            formatted_result = []
            speaker_ids = set()
            sentence_info = result[0].get("sentence_info", []) if result else []

            if not sentence_info:
                return Response({
                    "status": "success",
                    "filename": file.name,
                    "transcription": [{
                        "spk": "未知说话人",
                        "text": "未识别到有效内容",
                        "start_time": 0.00,
                        "end_time": 0.00,
                        "original_spk": "unknown"
                    }],
                    "speaker_stats": {"total_speakers": 0, "speaker_ids": [], "matched_speakers": {}},
                    "note": "已启用标点恢复+说话人识别+时间戳",
                    "timestamp": datetime.now().isoformat()
                })

            # ===================== 收集所有说话人的时间段 =====================
            from collections import defaultdict
            speaker_segments = defaultdict(list)
            for seg in sentence_info:
                spk_id = seg.get("spk") or seg.get("sp") or 0
                start_ms = seg.get("start", 0)
                end_ms = seg.get("end", 0)
                speaker_segments[spk_id].append((start_ms, end_ms))
                speaker_ids.add(spk_id)

            # ===================== 步骤2：对每个说话人，用ffmpeg切割 → 声纹匹配 =====================
            print(f"[{datetime.now()}] 第二步：说话人声纹匹配（共{len(speaker_segments)}人）")
            speaker_name_map = {}
            import subprocess

            for spk_id, segments in speaker_segments.items():
                # 选择最长的一段语音来识别（最准确）
                segments.sort(key=lambda x: x[1] - x[0], reverse=True)
                best_start, best_end = segments[0]

                start_sec = best_start / 1000.0
                end_sec = best_end / 1000.0
                duration_sec = end_sec - start_sec

                # 临时切割片段
                clip_path = os.path.join(settings.TEMP_DIR, f"speaker_clip_{spk_id}.wav")

                # ===================== ffmpeg 切割音频（核心！不需要pydub） =====================
                subprocess.run([
                    "ffmpeg",
                    "-ss", str(start_sec),
                    "-t", str(duration_sec),
                    "-i", temp_file,
                    "-ar", "16000",
                    "-ac", "1",
                    "-y",
                    clip_path
                ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

                # 提取声纹 + 匹配
                from .voiceprint_utils import extract_voiceprint_feature, match_voiceprint
                feat = extract_voiceprint_feature(clip_path)
                if feat is not None:
                    name = match_voiceprint(feat)
                    speaker_name_map[spk_id] = name if name else f"spk-{spk_id}"
                else:
                    speaker_name_map[spk_id] = f"spk-{spk_id}"

                # 删除临时片段
                if os.path.exists(clip_path):
                    os.remove(clip_path)

            # ===================== 最终结果替换名字 =====================
            matched_speakers = {}
            for spk_id, name in speaker_name_map.items():
                if not name.startswith("spk-"):
                    matched_speakers[spk_id] = name

            for seg in sentence_info:
                spk_id = seg.get("spk") or seg.get("sp") or 0
                name = speaker_name_map.get(spk_id, f"spk-{spk_id}")

                formatted_result.append({
                    "spk": name,
                    "text": seg.get("text", "").strip(),
                    "start_time": round(seg.get("start", 0) / 1000, 2),
                    "end_time": round(seg.get("end", 0) / 1000, 2),
                    "original_spk": spk_id
                })

            # ===================== 返回 =====================
            return Response({
                "status": "success",
                "filename": file.name,
                "transcription": formatted_result,
                "speaker_stats": {
                    "total_speakers": len(speaker_ids),
                    "speaker_ids": sorted(list(speaker_ids)),
                    "matched_speakers": matched_speakers
                },
                "note": "已启用标点恢复+说话人识别+时间戳 + 离线声纹匹配",
                "timestamp": datetime.now().isoformat()
            })

        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"处理失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

        finally:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass

# ------------------- 会议纪要生成接口 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class MeetingSummaryView(APIView):
    def post(self, request):
        """基于转录文本生成结构化会议纪要"""
        try:
            # 1. 提取请求参数
            transcription_text = request.data.get("transcription_text", "").strip()
            output_format = request.data.get("output_format", "txt")
            custom_system_prompt = request.data.get("custom_system_prompt", None)
            custom_user_prompt = request.data.get("custom_user_prompt", None)
            
            # 2. 输入校验
            if not transcription_text:
                return Response(
                    {"status": "failed", "detail": "转录文本不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 定义默认Prompt
            default_system_prompt = """你是专业的会议纪要生成助手，需严格按照以下要求处理：
1. 结构化提取信息：
   - 会议主题（精准概括核心议题）
   - 参与发言人列表（去重）
   - 核心讨论内容（按发言人分类整理，保留关键观点）
   - 关键结论/决议（明确会议达成的决策）
   - 待办事项（Action Items）：必须提取【任务-负责人】，若无截止时间则标注“未指定”，不要省略
2. 格式要求：
   - 使用层级标题（如## 会议主题）、项目符号/编号
   - 待办事项单独成节，突出显示
   - 语言简洁、逻辑清晰，无冗余信息
3. 特殊处理：
   - 多人对话严格区分发言人，避免混淆
   - 只要有任务，就必须列出来，不允许写“无待办事项”
   - 转录文本不完整时，基于已有内容合理归纳"""

            default_user_prompt = f"""请基于以下会议转录文本生成结构化会议纪要：

【会议转录文本】
{transcription_text}

【输出要求】
1. 严格遵循上述系统提示的结构和格式
2. 重点突出待办事项的"任务-负责人-截止时间"三元组
3. 仅输出纪要文本，无需额外解释或说明
4. 确保语言通顺，无语法错误，信息无遗漏"""
            
            # 4. 优先级：自定义 > 默认
            final_system_prompt = custom_system_prompt if (custom_system_prompt and custom_system_prompt.strip()) else default_system_prompt
            final_user_prompt = custom_user_prompt if (custom_user_prompt and custom_user_prompt.strip()) else default_user_prompt
            
            # 5. 调用LLM
            response = llm_client.chat.completions.create(
                model=settings.LLM_CONFIG["model_name"],
                messages=[
                    {"role": "system", "content": final_system_prompt},
                    {"role": "user", "content": final_user_prompt}
                ],
                stream=False,
                temperature=0.3,
                max_tokens=4000
            )
            
            # 6. 提取结果
            meeting_minutes = response.choices[0].message.content.strip()
            
            return Response({
                "status": "success",
                "meeting_minutes": meeting_minutes,
                "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S"),
                "output_format": output_format,
                "used_system_prompt": final_system_prompt,
                "used_user_prompt": final_user_prompt
            })
        
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"纪要生成失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- 会议摘要生成接口（轻量化） -------------------
@method_decorator(csrf_exempt, name='dispatch')
class MeetingAbstractView(APIView):
    def post(self, request):
        """基于转录文本生成轻量化会议摘要（核心要点提炼）"""
        try:
            # 1. 提取请求参数
            transcription_text = request.data.get("transcription_text", "").strip()
            abstract_length = request.data.get("abstract_length", "medium")  # short/medium/long
            custom_system_prompt = request.data.get("custom_system_prompt", None)
            custom_user_prompt = request.data.get("custom_user_prompt", None)
            
            # 2. 输入校验
            if not transcription_text:
                return Response(
                    {"status": "failed", "detail": "转录文本不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 定义长度对应的字数限制
            length_config = {
                "short": "100-200字",
                "medium": "200-400字",
                "long": "400-600字"
            }
            target_length = length_config.get(abstract_length, "200-400字")
            
            # 4. 定义默认Prompt（轻量化摘要）
            default_system_prompt = f"""你是专业的会议摘要生成助手，需严格按照以下要求处理：
1. 核心要求：
   - 仅提炼会议最核心的信息，拒绝冗余内容
   - 保留会议的核心议题、关键结论、重要待办
   - 语言高度凝练，符合{target_length}的字数要求
2. 格式要求：
   - 纯文本段落形式，无需层级标题和列表
   - 逻辑连贯，语句通顺，无语法错误
   - 不添加任何额外解释性文字"""

            default_user_prompt = f"""请基于以下会议转录文本生成轻量化会议摘要：

【会议转录文本】
{transcription_text}

【输出要求】
1. 严格遵循上述系统提示的字数要求（{target_length}）
2. 仅输出摘要文本，无需额外解释或说明
3. 确保覆盖核心议题、关键结论、重要待办
4. 语言简洁凝练，符合正式会议摘要的表达习惯"""
            
            # 5. 优先级：自定义 > 默认
            final_system_prompt = custom_system_prompt if (custom_system_prompt and custom_system_prompt.strip()) else default_system_prompt
            final_user_prompt = custom_user_prompt if (custom_user_prompt and custom_user_prompt.strip()) else default_user_prompt
            
            # 6. 调用LLM生成摘要
            response = llm_client.chat.completions.create(
                model=settings.LLM_CONFIG["model_name"],
                messages=[
                    {"role": "system", "content": final_system_prompt},
                    {"role": "user", "content": final_user_prompt}
                ],
                stream=False,
                temperature=0.2,  # 更低的温度保证摘要的准确性
                max_tokens=1000
            )
            
            # 7. 提取结果
            meeting_abstract = response.choices[0].message.content.strip()
            
            return Response({
                "status": "success",
                "meeting_abstract": meeting_abstract,
                "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S"),
                "abstract_length": abstract_length,
                "target_word_count": target_length,
                "used_system_prompt": final_system_prompt,
                "used_user_prompt": final_user_prompt
            })
        
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"摘要生成失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- Word导出接口（转录文本） -------------------
@method_decorator(csrf_exempt, name='dispatch')
class ExportTranscriptionWordView(APIView):
    def post(self, request):
        """导出语音分离结果为Word文档"""
        try:
            # 1. 提取参数
            transcription_text = request.data.get("transcription_text", "").strip()
            file_name = request.data.get("file_name", "会议记录")
            
            # 2. 输入校验
            if not transcription_text:
                return Response(
                    {"status": "failed", "detail": "语音分离文本不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 生成唯一文件名
            file_uuid = str(uuid.uuid4())
            doc_path = os.path.join(settings.TEMP_DIR, f"{file_name}_语音分离_{file_uuid}.docx")
            
            # 4. 创建Word文档
            doc = DocxDocument()
            
            # 标题
            title_para = doc.add_heading(level=0)
            title_run = title_para.add_run(f"{file_name} - 语音分离结果")
            title_run.font.size = Pt(20)
            title_run.bold = True
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 生成时间
            time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            time_para.add_run("\n\n")
            
            # 内容
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(transcription_text)
            content_run.font.size = Pt(12)
            content_run.font.name = "微软雅黑"
            content_para.line_spacing = 1.5
            
            # 保存文档
            doc.save(doc_path)
            
            # 5. 返回文件下载
            response = FileResponse(
                open(doc_path, 'rb'),
                filename=f"{file_name}_语音分离结果.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # 注：生产环境需添加文件清理逻辑（如定时任务）
            return response
        
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"生成语音分离Word失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- Word导出接口（会议纪要） -------------------
@method_decorator(csrf_exempt, name='dispatch')
class ExportSummaryWordView(APIView):
    def post(self, request):
        """导出会议纪要为Word文档"""
        try:
            # 1. 提取参数
            summary_text = request.data.get("summary_text", "").strip()
            file_name = request.data.get("file_name", "会议记录")
            
            # 2. 输入校验
            if not summary_text:
                return Response(
                    {"status": "failed", "detail": "会议纪要文本不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 生成唯一文件名
            file_uuid = str(uuid.uuid4())
            doc_path = os.path.join(settings.TEMP_DIR, f"{file_name}_会议纪要_{file_uuid}.docx")
            
            # 4. 创建Word文档
            doc = DocxDocument()
            
            # 标题
            title_para = doc.add_heading(level=0)
            title_run = title_para.add_run(f"{file_name} - 会议纪要")
            title_run.font.size = Pt(20)
            title_run.bold = True
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 生成时间
            time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            time_para.add_run("\n\n")
            
            # 内容
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(summary_text)
            content_run.font.size = Pt(12)
            content_run.font.name = "微软雅黑"
            content_para.line_spacing = 1.5
            
            # 保存文档
            doc.save(doc_path)
            
            # 5. 返回文件下载
            response = FileResponse(
                open(doc_path, 'rb'),
                filename=f"{file_name}_会议纪要.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            return response
        
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"生成会议纪要Word失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- Word导出接口（会议摘要） -------------------
@method_decorator(csrf_exempt, name='dispatch')
class ExportAbstractWordView(APIView):
    def post(self, request):
        """导出会议摘要为Word文档"""
        try:
            # 1. 提取参数
            abstract_text = request.data.get("abstract_text", "").strip()
            file_name = request.data.get("file_name", "会议记录")
            
            # 2. 输入校验
            if not abstract_text:
                return Response(
                    {"status": "failed", "detail": "会议摘要文本不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 生成唯一文件名
            file_uuid = str(uuid.uuid4())
            doc_path = os.path.join(settings.TEMP_DIR, f"{file_name}_会议摘要_{file_uuid}.docx")
            
            # 4. 创建Word文档
            doc = DocxDocument()
            
            # 标题
            title_para = doc.add_heading(level=0)
            title_run = title_para.add_run(f"{file_name} - 会议摘要")
            title_run.font.size = Pt(20)
            title_run.bold = True
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 生成时间
            time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            time_para.add_run("\n\n")
            
            # 摘要内容（特殊格式优化）
            abstract_para = doc.add_paragraph()
            abstract_run = abstract_para.add_run("【核心摘要】\n")
            abstract_run.bold = True
            abstract_run.font.size = Pt(14)
            
            content_run = abstract_para.add_run(abstract_text)
            content_run.font.size = Pt(12)
            content_run.font.name = "微软雅黑"
            abstract_para.line_spacing = 1.5
            
            # 保存文档
            doc.save(doc_path)
            
            # 5. 返回文件下载
            response = FileResponse(
                open(doc_path, 'rb'),
                filename=f"{file_name}_会议摘要.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            return response
        
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"生成会议摘要Word失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
# ------------------- 声纹添加 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class VoiceprintAddView(APIView):
    """添加声纹（上传音频提取特征，重复则提示）"""
    def post(self, request):
        try:
            # 1. 校验参数
            if 'file' not in request.FILES:
                return Response(
                    {"status": "failed", "detail": "未上传音频文件"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            voiceprint_name = request.POST.get("name", "").strip()
            if not voiceprint_name:
                return Response(
                    {"status": "failed", "detail": "声纹名称不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            # 校验名称是否已存在
            if Voiceprint.objects.filter(name=voiceprint_name).exists():
                return Response(
                    {"status": "failed", "detail": f"声纹名称「{voiceprint_name}」已存在"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 2. 校验音频格式
            file = request.FILES['file']
            if not file.name.lower().endswith(settings.ALLOWED_EXTENSIONS):
                return Response(
                    {"status": "failed", "detail": f"仅支持音频格式：{settings.ALLOWED_EXTENSIONS}"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 3. 保存临时音频文件
            temp_filename = f"vp_temp_{uuid.uuid4()}_{file.name}"
            temp_file = os.path.join(settings.TEMP_DIR, temp_filename)
            with open(temp_file, "wb") as f:
                for chunk in file.chunks():
                    f.write(chunk)
            
            # 4. 提取声纹特征
            global asr_model
            if asr_model is None:
                from .models import load_asr_model
                load_asr_model()
            vp_feature = extract_voiceprint_feature(temp_file)
            if vp_feature is None:
                return Response(
                    {"status": "failed", "detail": "声纹特征提取失败"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # 5. 检查是否重复
            is_duplicate, dup_name, max_sim = check_voiceprint_duplicate(vp_feature)
            if is_duplicate:
                return Response(
                    {"status": "failed", "detail": f"检测到相似声纹（相似度{max_sim:.2f}），名称：{dup_name}，不予添加"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 6. 保存声纹到数据库
            voiceprint = Voiceprint(
                name=voiceprint_name,
                feature=Voiceprint.feature_to_binary(vp_feature)
            )
            voiceprint.save()
            
            # 7. 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
            
            return Response({
                "status": "success",
                "detail": f"声纹「{voiceprint_name}」添加成功",
                "voiceprint_id": voiceprint.id,
                "created_at": voiceprint.created_at.isoformat()
            })
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"添加声纹失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- 声纹列表 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class VoiceprintListView(APIView):
    """获取声纹列表"""
    def get(self, request):
        try:
            voiceprints = Voiceprint.objects.all().order_by("-created_at")
            vp_list = [{
                "id": vp.id,
                "name": vp.name,
                "created_at": vp.created_at.isoformat(),
                "updated_at": vp.updated_at.isoformat()
            } for vp in voiceprints]
            return Response({
                "status": "success",
                "count": len(vp_list),
                "voiceprints": vp_list
            })
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"获取声纹列表失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- 声纹修改 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class VoiceprintRenameView(APIView):
    """修改声纹名称"""
    def post(self, request):
        try:
            # 1. 提取参数
            vp_id = request.data.get("id")
            new_name = request.data.get("new_name", "").strip()
            if not vp_id or not new_name:
                return Response(
                    {"status": "failed", "detail": "声纹ID和新名称不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 2. 检查声纹是否存在
            try:
                voiceprint = Voiceprint.objects.get(id=vp_id)
            except Voiceprint.DoesNotExist:
                return Response(
                    {"status": "failed", "detail": f"未找到ID为{vp_id}的声纹"},
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # 3. 检查新名称是否重复
            if Voiceprint.objects.filter(name=new_name).exclude(id=vp_id).exists():
                return Response(
                    {"status": "failed", "detail": f"声纹名称「{new_name}」已存在"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 4. 修改名称
            old_name = voiceprint.name
            voiceprint.name = new_name
            voiceprint.save()
            
            return Response({
                "status": "success",
                "detail": f"声纹名称从「{old_name}」修改为「{new_name}」成功",
                "voiceprint_id": vp_id
            })
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"修改声纹名称失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

# ------------------- 声纹删除 -------------------
@method_decorator(csrf_exempt, name='dispatch')
class VoiceprintDeleteView(APIView):
    """删除声纹"""
    def post(self, request):
        try:
            # 1. 提取参数
            vp_id = request.data.get("id")
            if not vp_id:
                return Response(
                    {"status": "failed", "detail": "声纹ID不能为空"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 2. 检查声纹是否存在
            try:
                voiceprint = Voiceprint.objects.get(id=vp_id)
            except Voiceprint.DoesNotExist:
                return Response(
                    {"status": "failed", "detail": f"未找到ID为{vp_id}的声纹"},
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # 3. 删除声纹
            vp_name = voiceprint.name
            voiceprint.delete()
            
            return Response({
                "status": "success",
                "detail": f"声纹「{vp_name}」删除成功",
                "voiceprint_id": vp_id
            })
        except Exception as e:
            traceback.print_exc()
            return Response(
                {"status": "failed", "detail": f"删除声纹失败：{str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )