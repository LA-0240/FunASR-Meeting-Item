from fastapi import FastAPI, File, UploadFile, HTTPException, Body
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from funasr import AutoModel
import os
import time
from openai import OpenAI
from datetime import datetime
from typing import Optional
import uuid
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# 初始化FastAPI应用
app = FastAPI(title="ASR会议纪要服务", version="1.0")

# 跨域配置（生产环境建议替换为具体域名）
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 配置项
LLM_CONFIG = {
    "api_key": "ms-8bc4a771-9cfc-42c9-956f-a725ee57aa09",  # 替换为实际的ModelScope Token
    "base_url": "https://api-inference.modelscope.cn/v1/",
    "model_name": "Qwen/Qwen3.5-35B-A3B"
}
TEMP_DIR = "./temp_files"  # 临时文件存储目录
ALLOWED_EXTENSIONS = ('.wav', '.mp3', '.ogg', '.flac')  # 支持的音频格式

# 创建临时目录
if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

# 初始化客户端/模型
llm_client = OpenAI(
    api_key=LLM_CONFIG["api_key"],
    base_url=LLM_CONFIG["base_url"]
)
asr_model = None

def load_asr_model():
    """加载ASR模型（含标点恢复、说话人识别、VAD）"""
    global asr_model
    if asr_model is None:
        print(f"{datetime.now()} - 开始加载ASR模型（CPU模式）...")
        try:
            asr_model = AutoModel(
                model="paraformer-zh",
                vad_model="fsmn-vad",        # 语音活动检测
                punc_model="ct-punc",        # 标点恢复
                spk_model="cam++",           # 说话人识别
                device="cpu"                 # 生产环境可改为"cuda"
            )
            print(f"{datetime.now()} - ASR模型加载完成")
        except Exception as e:
            print(f"{datetime.now()} - ASR模型加载失败: {str(e)}")
            raise

# 应用启动时加载模型
@app.on_event("startup")
async def startup_event():
    load_asr_model()

# ------------------- 基础接口 -------------------
@app.get("/", tags=["基础接口"])
async def health_check():
    """服务健康检查"""
    return {
        "status": "healthy",
        "service": "ASR会议纪要服务",
        "timestamp": datetime.now().isoformat(),
        "features": ["语音转文字（带标点/说话人）", "会议纪要生成", "Word文档导出"]
    }

# ------------------- ASR语音转文字接口 -------------------
@app.post("/asr", tags=["语音转文字"])
async def transcribe_audio(
    file: UploadFile = File(..., description="音频文件（支持wav/mp3/ogg/flac）"),
    batch_size_s: int = Body(300, description="批处理时长（秒）"),
    hotword: Optional[str] = Body(None, description="热词（提升特定词汇识别准确率）")
):
    """
    语音转文字，支持：
    - 多格式音频文件
    - 标点自动恢复
    - 说话人识别与分段
    - VAD语音活动检测
    """
    temp_file = None
    try:
        # 1. 文件类型校验
        if not file.filename.lower().endswith(ALLOWED_EXTENSIONS):
            raise HTTPException(
                status_code=400,
                detail=f"仅支持以下音频格式：{ALLOWED_EXTENSIONS}"
            )

        # 2. 保存临时文件
        temp_file = os.path.join(TEMP_DIR, f"temp_{datetime.now().strftime('%Y%m%d%H%M%S_%f')}_{file.filename}")
        with open(temp_file, "wb") as buffer:
            buffer.write(await file.read())

        # 3. 调用ASR模型处理音频
        print(f"{datetime.now()} - 开始处理音频文件: {file.filename}")
        result = asr_model.generate(
            input=temp_file,
            batch_size_s=batch_size_s,
            hotword=hotword,
            punc=True,                # 启用标点恢复
            spk_segment=True,         # 启用说话人分段
            merge_vad=True,           # 合并同说话人连续片段
            max_single_segment_time=30 # 单个说话人片段最大时长
        )

        # 4. 格式化识别结果
        formatted_result = []
        if result and len(result) > 0:
            # 优先读取sentence_info字段（新版funasr返回格式）
            sentence_info = result[0].get("sentence_info", [])
            if sentence_info:
                for segment in sentence_info:
                    formatted_result.append({
                        "spk": segment.get("spk", "未知说话人"),
                        "text": segment.get("text", "").strip()
                    })
            # 兼容旧版value字段
            elif "value" in result[0]:
                for segment in result[0]["value"]:
                    formatted_result.append({
                        "spk": segment.get("spk", "未知说话人"),
                        "text": segment.get("text", "").strip()
                    })
        
        # 处理空结果
        if not formatted_result:
            formatted_result = [{"spk": "未知说话人", "text": "未识别到有效语音内容"}]

        # 5. 返回结果
        return JSONResponse(content={
            "status": "success",
            "filename": file.filename,
            "transcription": formatted_result,
            "note": "已启用标点恢复+说话人识别（CPU模式）",
            "timestamp": datetime.now().isoformat()
        })

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"音频处理失败：{str(e)}")
    finally:
        # 清理临时文件
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
                print(f"{datetime.now()} - 临时文件 {temp_file} 已删除")
            except Exception as e:
                print(f"{datetime.now()} - 临时文件删除失败：{str(e)}")

# ------------------- 会议纪要生成接口 -------------------
@app.post("/generate_summary", tags=["会议纪要"])
async def generate_summary(
    transcription_text: str = Body(..., description="会议转录文本（带说话人标识）"),
    output_format: str = Body("txt", description="输出格式（暂支持txt）")
):
    """
    基于会议转录文本生成结构化会议纪要
    - 提取会议主题、参与人、讨论内容、决议、待办事项等
    - 待办事项需明确 任务-负责人-截止时间
    """
    try:
        # 输入校验
        if not transcription_text or transcription_text.strip() == "":
            raise HTTPException(status_code=400, detail="转录文本不能为空")
        
        # 系统提示词（定义纪要生成规则）
        system_prompt = """你是专业的会议纪要生成助手，需严格按照以下要求处理：
1. 结构化提取信息：
   - 会议主题（精准概括核心议题）
   - 参与发言人列表（去重）
   - 核心讨论内容（按发言人分类整理，保留关键观点）
   - 关键结论/决议（明确会议达成的决策）
   - 待办事项（Action Items）：必须明确【任务-负责人-截止时间】，无截止时间标注"未指定"
2. 格式要求：
   - 使用层级标题（如## 会议主题）、项目符号/编号
   - 待办事项单独成节，突出显示
   - 语言简洁、逻辑清晰，无冗余信息
3. 特殊处理：
   - 多人对话严格区分发言人，避免混淆
   - 无待办事项时明确标注"无待办事项"
   - 转录文本不完整时，基于已有内容合理归纳"""

        # 用户提示词（传入具体转录文本）
        user_prompt = f"""请基于以下会议转录文本生成结构化会议纪要：

【会议转录文本】
{transcription_text}

【输出要求】
1. 严格遵循上述系统提示的结构和格式
2. 重点突出待办事项的"任务-负责人-截止时间"三元组
3. 仅输出纪要文本，无需额外解释或说明
4. 确保语言通顺，无语法错误，信息无遗漏"""

        # 调用LLM生成纪要
        response = llm_client.chat.completions.create(
            model=LLM_CONFIG["model_name"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            stream=False,
            temperature=0.3,  # 低温度保证输出稳定
            max_tokens=4000   # 足够容纳长纪要
        )

        # 提取并处理结果
        meeting_minutes = response.choices[0].message.content.strip()
        
        return {
            "status": "success",
            "meeting_minutes": meeting_minutes,
            "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S"),
            "output_format": output_format
        }

    except HTTPException:
        raise  # 直接抛出已定义的HTTP异常
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"纪要生成失败：{str(e)}")

# ------------------- 新增：Word导出接口 -------------------
@app.post("/export_transcription_word", tags=["Word导出"])
async def export_transcription_word(
    transcription_text: str = Body(..., description="语音分离文本（带说话人标识）"),
    file_name: str = Body("会议记录", description="文件名前缀")
):
    """导出语音分离结果为Word文档"""
    try:
        # 输入校验
        if not transcription_text or transcription_text.strip() == "":
            raise HTTPException(status_code=400, detail="语音分离文本不能为空")
        
        # 生成唯一文件名（避免冲突）
        file_uuid = str(uuid.uuid4())
        doc_path = os.path.join(TEMP_DIR, f"{file_name}_语音分离_{file_uuid}.docx")

        # 1. 创建Word文档
        doc = DocxDocument()
        
        # 2. 设置文档标题
        title_para = doc.add_heading(level=0)
        title_run = title_para.add_run(f"{file_name} - 语音分离结果")
        title_run.font.size = Pt(20)
        title_run.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 3. 添加生成时间
        time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        time_para.add_run("\n\n")

        # 4. 添加语音分离内容（按行拆分，保持格式）
        content_para = doc.add_paragraph()
        content_run = content_para.add_run(transcription_text)
        content_run.font.size = Pt(12)
        content_run.font.name = "微软雅黑"
        content_para.line_spacing = 1.5  # 1.5倍行间距

        # 5. 保存文档
        doc.save(doc_path)

        # 6. 返回文件供下载
        return FileResponse(
            path=doc_path,
            filename=f"{file_name}_语音分离结果.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成语音分离Word失败：{str(e)}")

@app.post("/export_summary_word", tags=["Word导出"])
async def export_summary_word(
    summary_text: str = Body(..., description="会议纪要文本"),
    file_name: str = Body("会议记录", description="文件名前缀")
):
    """导出会议纪要为Word文档"""
    try:
        # 输入校验
        if not summary_text or summary_text.strip() == "":
            raise HTTPException(status_code=400, detail="会议纪要文本不能为空")
        
        # 生成唯一文件名
        file_uuid = str(uuid.uuid4())
        doc_path = os.path.join(TEMP_DIR, f"{file_name}_会议纪要_{file_uuid}.docx")

        # 1. 创建Word文档
        doc = DocxDocument()
        
        # 2. 设置文档标题
        title_para = doc.add_heading(level=0)
        title_run = title_para.add_run(f"{file_name} - 会议纪要")
        title_run.font.size = Pt(20)
        title_run.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 3. 添加生成时间
        time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        time_para.add_run("\n\n")

        # 4. 添加会议纪要内容
        content_para = doc.add_paragraph()
        content_run = content_para.add_run(summary_text)
        content_run.font.size = Pt(12)
        content_run.font.name = "微软雅黑"
        content_para.line_spacing = 1.5

        # 5. 保存文档
        doc.save(doc_path)

        # 6. 返回文件供下载
        return FileResponse(
            path=doc_path,
            filename=f"{file_name}_会议纪要.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成会议纪要Word失败：{str(e)}")

# ------------------- 主程序入口 -------------------
if __name__ == "__main__":
    import uvicorn
    # CPU模式建议workers=1，GPU模式可适当增加
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        workers=1,
        log_level="info"
    )