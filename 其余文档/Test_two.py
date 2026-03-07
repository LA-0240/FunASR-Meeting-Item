import requests
import time
from datetime import datetime
from openai import OpenAI
import docx  # 需安装：pip install python-docx

def test_asr_service(audio_file_path, hotword=None):
    """调用本地ASR服务，识别音频并返回结构化转录文本（含说话人）"""
    # API端点
    url = "http://localhost:8000/asr"

    # 准备文件
    with open(audio_file_path, "rb") as audio_file:
        files = {"file": (audio_file_path.split("/")[-1], audio_file)}
        data = {"hotword": hotword} if hotword else {}

        # 发送请求
        start_time = time.time()
        print(f"{datetime.now()} - 开始发送ASR请求...")
        response = requests.post(url, files=files, data=data)

        # 处理响应
        if response.status_code == 200:
            result = response.json()
            print(f"{datetime.now()} - ASR识别成功!")
            print(f"文件名: {result['filename']}")
            print("识别结果（带标点+说话人）：")
            
            # 拼接结构化转录文本（用于LLM处理）
            transcription_text = ""
            for idx, segment in enumerate(result['transcription']):
                spk = segment['spk']
                text = segment['text']
                print(f"\n片段{idx+1}：")
                print(f"说话人：{spk}")
                print(f"文本：{text}")
                transcription_text += f"【说话人{spk}】：{text}\n"
            
            print(f"\nASR总处理时间: {time.time() - start_time:.2f}秒")
            return result, transcription_text
        else:
            print(f"{datetime.now()} - ASR识别失败!")
            print(f"状态码: {response.status_code}")
            print(f"错误信息: {response.text}")
            return None, None

def generate_meeting_minutes(transcription_text, output_format="txt"):
    """调用ModelScope LLM生成结构化会议纪要，并导出为指定格式"""
    if not transcription_text:
        print("错误：转录文本为空，无法生成会议纪要")
        return None

    # 初始化ModelScope OpenAI兼容客户端
    client = OpenAI(
        api_key="ms-8bc4a771-9cfc-42c9-956f-a725ee57aa09",  # 替换为你的ModelScope Access Token
        base_url="https://api-inference.modelscope.cn/v1/"
    )

    # 定义会议纪要生成的Prompt（按要求定制）
    system_prompt = """你是专业的会议纪要生成助手，需完成以下任务：
1. 基于音频转录文本，提取结构化会议信息，包含：
   - 会议主题
   - 参与发言人列表
   - 核心讨论内容（分发言人整理）
   - 关键结论/决议
   - 待办事项（Action Items，需明确【任务-负责人-截止时间】，无截止时间则标注“未指定”）
2. 输出格式要求：
   - 层级清晰，使用标题、分点（序号/项目符号）排版
   - 待办事项单独标注，突出“任务-负责人-截止时间”
   - 语言简洁、逻辑连贯，保留关键信息不遗漏
3. 处理多人对话时，严格区分各说话人的发言内容，避免混淆"""

    user_prompt = f"""请基于以下会议音频转录文本，生成结构化会议纪要：

【会议转录文本】
{transcription_text}

【输出要求】
- 严格按照上述系统提示的结构生成
- 重点提取“任务-负责人-截止时间”类结构化信息
- 最终输出完整的会议纪要文本，无需额外说明"""

    try:
        print("\n开始调用ModelScope LLM生成会议纪要...")
        start_time = time.time()
        
        # 调用LLM（流式返回）
        response = client.chat.completions.create(
            model="Qwen/Qwen3.5-35B-A3B",  # ModelScope Model-Id
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            stream=True,
            temperature=0.3,  # 降低随机性，保证输出结构化
            max_tokens=4000
        )

        # 拼接流式返回结果
        meeting_minutes = ""
        print("\n===== 生成的会议纪要 =====")
        for chunk in response:
            content = chunk.choices[0].delta.content
            if content:
                meeting_minutes += content
                print(content, end='', flush=True)
        
        print(f"\n\nLLM处理时间: {time.time() - start_time:.2f}秒")

        # 导出为指定格式（TXT/Word）
        if meeting_minutes:
            file_name = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            if output_format.lower() == "txt":
                file_path = f"{file_name}.txt"
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(meeting_minutes)
                print(f"\n会议纪要已导出为TXT文件：{file_path}")
            elif output_format.lower() == "docx":
                file_path = f"{file_name}.docx"
                doc = docx.Document()
                doc.add_heading('会议纪要', 0)
                doc.add_paragraph(meeting_minutes)
                doc.save(file_path)
                print(f"\n会议纪要已导出为Word文件：{file_path}")
            return meeting_minutes
        else:
            print("错误：LLM返回空的会议纪要")
            return None

    except Exception as e:
        print(f"\nLLM调用/纪要生成失败：{str(e)}")
        return None

if __name__ == "__main__":
    # 配置参数
    audio_file = r"F:\Text_FunASR\FunASR-main\TTT.mp3"  # 替换为你的音频文件路径
    hotword = "创业项目,宠物项圈,VR非遗保护"  # 可选热词
    output_format = "docx"  # 输出格式：txt / docx

    # 1. 调用ASR识别音频，获取转录文本（含说话人）
    asr_result, transcription_text = test_asr_service(audio_file, hotword)
    
    # 2. 调用LLM生成并导出结构化会议纪要
    if asr_result and transcription_text:
        meeting_minutes = generate_meeting_minutes(transcription_text, output_format)
    else:
        print("ASR识别失败，无法生成会议纪要")